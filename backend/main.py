"""
Писарь v3 — Backend API
FastAPI + Nexara STT + Anthropic Claude + SQLite + Auth
"""

import os
import json
import re
import tempfile
import sqlite3
import uuid
import httpx
import hashlib
import secrets
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
import warnings
warnings.filterwarnings("ignore", message="Unverified HTTPS request")

app = FastAPI(title="Писарь API", version="3.0.0")

# Serve React static files in production
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import pathlib

STATIC_DIR = pathlib.Path(__file__).parent / "static"
DB_PATH = pathlib.Path(__file__).parent / "pisar.db"

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Database ───

def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            name TEXT DEFAULT '',
            login TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            token TEXT UNIQUE,
            created_at TEXT DEFAULT ''
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS records (
            id TEXT PRIMARY KEY,
            user_id TEXT DEFAULT '',
            patient_name TEXT DEFAULT '',
            diagnosis_code TEXT DEFAULT '',
            specialty TEXT DEFAULT '',
            summary TEXT DEFAULT '',
            sections TEXT DEFAULT '[]',
            transcript TEXT DEFAULT '',
            created_at TEXT DEFAULT ''
        )
    """)
    # Миграция: добавить user_id если его нет (для старых баз)
    try:
        conn.execute("ALTER TABLE records ADD COLUMN user_id TEXT DEFAULT ''")
    except Exception:
        pass  # Колонка уже существует
    conn.commit()
    conn.close()


init_db()

# GigaChat API
GIGACHAT_AUTH_KEY = os.environ.get("GIGACHAT_AUTH_KEY", "")
GIGACHAT_TOKEN_URL = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
GIGACHAT_API_URL = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"
GIGACHAT_MODEL = os.environ.get("GIGACHAT_MODEL", "GigaChat-Pro")


async def gigachat_complete(messages: list, max_tokens: int = 8192) -> str:
    """Выполняет запрос к GigaChat API и возвращает текст ответа."""
    async with httpx.AsyncClient(verify=False) as client:
        token_resp = await client.post(
            GIGACHAT_TOKEN_URL,
            headers={
                "Authorization": f"Basic {GIGACHAT_AUTH_KEY}",
                "RqUID": str(uuid.uuid4()),
                "Content-Type": "application/x-www-form-urlencoded",
            },
            data={"scope": "GIGACHAT_API_PERS"},
            timeout=30.0,
        )
        token_resp.raise_for_status()
        access_token = token_resp.json()["access_token"]

    async with httpx.AsyncClient(verify=False) as client:
        resp = await client.post(
            GIGACHAT_API_URL,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json={
                "model": GIGACHAT_MODEL,
                "messages": messages,
                "max_tokens": max_tokens,
            },
            timeout=120.0,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]

# Nexara API для распознавания речи
NEXARA_API_URL = "https://api.nexara.ru/api/v1/audio/transcriptions"
NEXARA_API_KEY = os.environ.get("NEXARA_API_KEY", "")

ALLOWED_AUDIO = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".mp4", ".mpeg", ".mpga", ".oga", ".wma", ".aac"}


# ─── Auth helpers ───

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()


def get_user_by_token(token: str):
    if not token:
        return None
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE token = ?", (token,)).fetchone()
    conn.close()
    return dict(row) if row else None


def require_auth(authorization: str = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="Необходима авторизация")
    token = authorization.replace("Bearer ", "")
    user = get_user_by_token(token)
    if not user:
        raise HTTPException(status_code=401, detail="Неверный токен")
    return user


# ─── Auth endpoints ───

@app.post("/auth/register")
async def register(login: str = Form(...), password: str = Form(...), name: str = Form("")):
    if len(login) < 3:
        raise HTTPException(status_code=400, detail="Логин должен быть минимум 3 символа")
    if len(password) < 4:
        raise HTTPException(status_code=400, detail="Пароль должен быть минимум 4 символа")
    conn = get_db()
    existing = conn.execute("SELECT id FROM users WHERE login = ?", (login,)).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=400, detail="Пользователь с таким логином уже существует")
    user_id = str(uuid.uuid4())[:8]
    token = secrets.token_hex(32)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    conn.execute(
        "INSERT INTO users (id, name, login, password_hash, token, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (user_id, name or login, login, hash_password(password), token, now),
    )
    conn.commit()
    conn.close()
    return {"token": token, "user": {"id": user_id, "name": name or login, "login": login}}


@app.post("/auth/login")
async def auth_login(login: str = Form(...), password: str = Form(...)):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE login = ?", (login,)).fetchone()
    conn.close()
    if not row or row["password_hash"] != hash_password(password):
        raise HTTPException(status_code=401, detail="Неверный логин или пароль")
    user = dict(row)
    # Обновляем токен при каждом входе
    token = secrets.token_hex(32)
    conn = get_db()
    conn.execute("UPDATE users SET token = ? WHERE id = ?", (token, user["id"]))
    conn.commit()
    conn.close()
    return {"token": token, "user": {"id": user["id"], "name": user["name"], "login": user["login"]}}


@app.get("/auth/me")
async def auth_me(authorization: str = Header(None)):
    user = require_auth(authorization)
    return {"id": user["id"], "name": user["name"], "login": user["login"]}


# ─── Расширенные промпты по специальностям ───

PROMPTS = {
    "psychiatrist": """Ты — ИИ-ассистент психиатра психоневрологического диспансера (ПНД). Получив расшифровку речи врача, структурируй её в документ первичного осмотра пациента по ТОЧНОМУ формату ПНД.

Формат ответа — СТРОГО JSON (без markdown, без backticks):
{
  "patient_name": "Фамилия пациента если упомянута, иначе пустая строка",
  "date": "Дата приёма если упомянута, иначе пустая строка",
  "specialty": "Психиатр",
  "diagnosis_code": "Код МКБ-10 если определён, иначе пустая строка",
  "sections": [
    {
      "title": "Обращение",
      "content": "Обратился в ДС: первично/повторно; цель (для обследования и подбора терапии / для лечения / для обследования по линии РВК и т.д.)."
    },
    {
      "title": "Жалобы",
      "content": "Жалобы пациента в кавычках, как высказывает сам пациент. Если активно не предъявляет — указать."
    },
    {
      "title": "Анамнез жизни",
      "content": "Место рождения, состав семьи (полная/неполная, братья/сёстры). Данные о психопатологически отягощённой наследственности. Раннее развитие. ДДУ. Школа (возраст, успеваемость, отношения со сверстниками, буллинг). Дополнительное образование, кружки, секции. Дальнейшее образование (ВУЗ, ПТУ, специальность). Трудовая деятельность (где работал, текущее трудоустройство). Семейное положение, романтические отношения, дети. Проживание (с кем, условия). Лицензия на ношение оружия (имеет/не имеет). Водительские права (имеет/не имеет)."
    },
    {
      "title": "Криминальный анамнез",
      "content": "Судим/не судим. Привлекался ли к административной и уголовной ответственности."
    },
    {
      "title": "Перенесённые заболевания",
      "content": "Детские инфекции, ОРВИ. ЧМТ (когда, сколько раз, стационарное лечение). Операции (какие, когда). Хронические соматические заболевания. Другие травмы, переливания крови, судорожные припадки."
    },
    {
      "title": "Данные обследований",
      "content": "Результаты имеющихся обследований с датами и источниками: рентгенография, ЭКГ, анализы крови (БАК, КАК, HBsAg, AntiHCV), осмотры специалистов, ЭПО и др. Если данных нет — 'Актуальных данных обследований не предоставил'."
    },
    {
      "title": "Аллергоанамнез",
      "content": "Аллергические реакции на лекарства, продукты и др. с указанием характера реакции (отёк, сыпь и т.д.). Если нет — 'спокойный'."
    },
    {
      "title": "Эпиданамнез",
      "content": "Отягощён / не отягощён. Контакты с инфекционными больными за последние 14 дней."
    },
    {
      "title": "Наркоанамнез",
      "content": "Курение (курит/не курит, количество). Алкоголь (употребляет/не употребляет, частота, количество). Наркотические вещества (употребление, вид, частота, формирование зависимости). Если отрицает — указать."
    },
    {
      "title": "Анамнез заболевания",
      "content": "Данные за психопатологически отягощённую наследственность. Преморбидные особенности личности (характер с детства, общительность, замкнутость). Дебют психических нарушений (когда, при каких обстоятельствах, провоцирующие факторы). Хронологическое описание клинической картины: симптомы, их динамика, периоды ухудшения и улучшения. Обращения за психиатрической помощью (когда, куда, стационарно/амбулаторно). Проводившееся лечение (препараты с дозировками, эффективность, побочные эффекты, причины смены терапии). Выписные диагнозы. Текущее лечение и его эффективность. Причина текущей госпитализации/обращения."
    },
    {
      "title": "Психическое состояние",
      "content": "Сознание (формально не помрачено / помрачено). Ориентировка в месте, времени и собственной личности. Внешний вид (опрятен, неопрятен, особенности). Контакт (охотно/неохотно вступает в беседу, формально). Поведение (тревожен, суетлив, спокоен, заторможен). Речь (темп, спонтанность, многоречивость/малоговорящий). Мышление (темп, структурные нарушения — разноплановость, детализация, резонёрство, обстоятельность; содержание — навязчивые, сверхценные, бредовые идеи). Обманы восприятия (галлюцинации — отрицает/обнаруживает). Эмоциональная сфера (фон настроения, аффект, адекватность реакций, ангедония). Волевая сфера (мотивация, активность). Агрессивные и аутоагрессивные тенденции. Внимание и память. Интеллект (соответствие возрасту и образованию). Сон (продолжительность, нарушения). Аппетит. Критика к своему состоянию (полная / неполная / отсутствует)."
    },
    {
      "title": "Соматический статус",
      "content": "Общее состояние. Рост, вес, ИМТ, обхват талии. Кожные покровы и видимые слизистые. Задняя стенка глотки. Периферические лимфоузлы. Сердечные тоны (ясные/приглушённые, ритмичные, шумы). АД, ЧСС. Дыхание (везикулярное, хрипы). Живот (мягкий, безболезненный). Печень, селезёнка. Симптом поколачивания. Периферические отёки. Физиологические отправления."
    },
    {
      "title": "Неврологический статус",
      "content": "Лицо (симметрично/асимметрично). Зрачки (равновеликие, D=S, фотореакции). Язык (по средней линии/девиация). Глотание и фонация. Сухожильные рефлексы (D=S, живые/снижены). Мышечный тонус. Нарушения чувствительности, парезы, параличи. Поза Ромберга. Пальценосовая проба. Очаговая и менингеальная симптоматика. Пальпация остистых отростков."
    },
    {
      "title": "Обоснование диагноза",
      "content": "Краткое обоснование диагноза на основании данных анамнеза, клинической картины, динамики состояния и ответа на терапию. Дифференциальная диагностика при необходимости."
    },
    {
      "title": "Диагноз",
      "content": "Основной диагноз в формате: синдром, состояние ремиссии/обострения, код МКБ-10. Сопутствующие диагнозы."
    },
    {
      "title": "Социальный статус",
      "content": "Учёба/работа, инвалидность, состоит ли в ЦЗН."
    },
    {
      "title": "План обследования и лечения",
      "content": "Режим. Диета. Плановые обследования (консультации специалистов, ЭПО, ЭЭГ, анализы, направления). Текущая фармакотерапия (препарат, дозировка, время приёма, цель назначения). Планируемые изменения терапии. Психотерапия (индивидуальная, групповая). Коррекция по состоянию."
    }
  ],
  "summary": "Краткое резюме осмотра в 1-2 предложения"
}

Правила:
- Используй ТОЧНЫЙ стиль написания психиатрических записей ПНД (как в приведённых примерах эталонных записей)
- Жалобы пиши в кавычках — как говорит сам пациент
- Анамнез жизни — связным текстом, НЕ списком
- Анамнез заболевания — хронологически, связным текстом с конкретными датами и описанием клинической картины
- Психическое состояние — связным текстом, описательно
- Соматический и неврологический статус — кратко, по пунктам через точку
- Обоснование диагноза — аргументированный связный текст
- Заполняй разделы ТОЛЬКО на основе предоставленных данных
- Если данных для раздела нет — напиши "Данные не предоставлены"
- НЕ придумывай информацию, которой нет в расшифровке
- Пиши на русском языке
- Формальную преамбулу (представился психиатром, разъяснены права, согласие на осмотр) НЕ включай — она добавляется автоматически""",

    "psychiatrist_diary": """Ты — ИИ-ассистент психиатра ПНД. На основе предоставленного анамнеза, диагноза, терапии и описания состояния пациента составь серию дневниковых записей наблюдения.

ВАЖНО: В тексте врача будет указан период генерации (1 неделя / 2 недели / 1 месяц). Создай записи с датами через 2-4 дня друг от друга на весь указанный период. Дат должно быть:
- 1 неделя = 2-3 записи
- 2 недели = 4-5 записей
- 1 месяц = 8-10 записей

Каждая запись — это отдельная дата наблюдения. Записи должны показывать ДИНАМИКУ состояния: улучшение, ухудшение, стабилизацию, побочные эффекты терапии, изменения сна, настроения, тревоги.

Формат ответа — СТРОГО JSON (без markdown, без backticks):
{
  "patient_name": "Фамилия пациента если упомянута",
  "date": "",
  "specialty": "Психиатр (дневник)",
  "diagnosis_code": "Код МКБ-10 если определён",
  "sections": [
    {
      "title": "ДД.ММ.ГГГГ",
      "content": "Краткая запись наблюдения за этот день. Включает: жалобы пациента (в его словах), сон, аппетит, настроение, тревога, побочные эффекты терапии, поведение, социальное функционирование. 2-5 предложений."
    },
    {
      "title": "ДД.ММ.ГГГГ",
      "content": "Следующая запись с динамикой..."
    }
  ],
  "summary": "Общая динамика за период в 1 предложение"
}

Примеры стиля записей (ориентируйся на этот формат):

"03.03.2026
Раздражительность есть, тревожность небольшая. Сон в порядке. Сильная зажатость, ощущение что все смотрят."

"05.02.2026
Была сонливость, но вечером не было агрессии. Колебания настроения не выражены. Фантазии в последнее время вернулись."

"11.02.2026
Гораздо более собран, легче заниматься делами. Сон без пробуждений, 7-8 часов. Настроение пока не появляется, сейчас гораздо легче находиться на людях, меньше напряжения. Эмоции вновь чувствую."

"21.01.2026
Спокоен, ровен. Сегодня тревоги не было, чувствую себя стабильно. Ложусь в 11, ночью не просыпаюсь, сон глубокий."

"13.01.2026
Сознание формально не помрачено. Ориентирован в месте, времени и собственной личности верно. Внешне представляется тревожным, суетливым. Охотно вступает в беседу. Мышление без грубых структурных нарушений. Фон настроения определяется лекарственной эутимией с тревожным компонентом. Сон без пробуждений, 7-8 часов. Аппетит достаточный. Критика к своему состоянию неполная."

Правила:
- Записи должны быть КРАТКИМИ — 2-5 предложений каждая, как в примерах выше
- Чередуй стили: иногда от лица пациента (его слова), иногда описание врача
- Показывай реалистичную динамику: не всё улучшается линейно, бывают откаты
- Упоминай побочные эффекты терапии (сонливость, набор веса и т.д.) если релевантно
- Используй даты начиная от сегодняшней или указанной врачом
- Опирайся на диагноз, терапию и анамнез пациента
- Пиши на русском языке""",

    "therapist": """Ты — ИИ-ассистент терапевта. Получив расшифровку речи врача, структурируй её в полноценный медицинский документ терапевтического приёма.

Формат ответа — СТРОГО JSON (без markdown, без backticks):
{
  "patient_name": "ФИО пациента если упомянуто, иначе пустая строка",
  "date": "Дата приёма если упомянута, иначе пустая строка",
  "specialty": "Терапевт",
  "diagnosis_code": "Код МКБ-10 если определён, иначе пустая строка",
  "sections": [
    {
      "title": "Жалобы",
      "content": "Основные жалобы с детализацией каждой:\n- Локализация, иррадиация\n- Характер (жжение, покалывание, давление и т.д.)\n- Провоцирующие факторы\n- Продолжительность и интенсивность\n- Что приносит облегчение\n- Сопутствующие симптомы\n- Дополнительные жалобы"
    },
    {
      "title": "Анамнез заболевания (Anamnesis morbi)",
      "content": "Хронологическое описание:\n- Когда и при каких обстоятельствах заболел\n- Первые проявления\n- Обращения к врачам, обследования, диагнозы\n- Проводившееся лечение и его эффективность\n- Динамика симптомов\n- Частота обострений\n- Трудоспособность"
    },
    {
      "title": "Анамнез жизни (Anamnesis vitae)",
      "content": "- Биографические данные\n- Семейно-половой анамнез\n- Трудовой анамнез, профвредности\n- Бытовые условия, питание\n- Вредные привычки\n- Перенесённые заболевания\n- Эпидемиологический анамнез\n- Аллергологический анамнез\n- Наследственность"
    },
    {
      "title": "Объективный осмотр (Status praesens)",
      "content": "- Общее состояние, сознание, положение\n- Телосложение, рост, вес, ИМТ\n- Кожные покровы и слизистые\n- Лимфатические узлы\n- Опорно-двигательная система\n- Органы дыхания (ЧДД, перкуссия, аускультация)\n- Сердечно-сосудистая система (АД, ЧСС, тоны, шумы)\n- Органы пищеварения (язык, живот, печень, селезёнка)\n- Мочевыделительная система\n- Нервная система\n- Температура тела"
    },
    {
      "title": "Диагноз",
      "content": "Основной диагноз по МКБ-10.\nСопутствующие заболевания."
    },
    {
      "title": "Назначения",
      "content": "- Медикаментозная терапия (препарат, дозировка, кратность, длительность)\n- Немедикаментозное лечение\n- Обследования (анализы, ЭКГ, УЗИ и т.д.)\n- Консультации специалистов\n- Дата повторного приёма\n- Рекомендации по режиму и питанию"
    }
  ],
  "summary": "Краткое резюме приёма в 1-2 предложения"
}

Правила:
- Профессиональная медицинская терминология
- Заполняй ТОЛЬКО на основе предоставленных данных
- Нет данных = "Данные не предоставлены"
- НЕ придумывай информацию
- МКБ-10, русский язык""",

    "pediatrician": """Ты — ИИ-ассистент педиатра. Получив расшифровку речи врача, структурируй её в полноценный медицинский документ педиатрического приёма.

Формат ответа — СТРОГО JSON (без markdown, без backticks):
{
  "patient_name": "ФИО ребёнка если упомянуто, иначе пустая строка",
  "date": "Дата приёма если упомянута, иначе пустая строка",
  "specialty": "Педиатр",
  "diagnosis_code": "Код МКБ-10 если определён, иначе пустая строка",
  "sections": [
    {
      "title": "Жалобы",
      "content": "Жалобы (со слов родителей/ребёнка):\n- Основные жалобы с детализацией\n- Дополнительные жалобы"
    },
    {
      "title": "Анамнез заболевания",
      "content": "- Когда заболел, с чего началось\n- Динамика симптомов\n- Проводившееся лечение до обращения\n- Эффективность лечения"
    },
    {
      "title": "Анамнез жизни",
      "content": "- Течение беременности и родов у матери\n- Вес и рост при рождении, оценка по Апгар\n- Вскармливание (грудное/искусственное)\n- Перенесённые заболевания\n- Прививки (по календарю/нет)\n- Аллергологический анамнез\n- Наследственность"
    },
    {
      "title": "Объективный осмотр",
      "content": "- Общее состояние, сознание, активность\n- Температура тела\n- Кожные покровы и слизистые\n- Зев, миндалины\n- Лимфатические узлы\n- Носовое дыхание\n- Органы дыхания (ЧДД, аускультация)\n- Сердечно-сосудистая система (ЧСС)\n- Живот\n- Стул, мочеиспускание"
    },
    {
      "title": "Физическое развитие",
      "content": "- Возраст\n- Вес, рост\n- Соответствие возрастным нормам\n- Центильные коридоры если возможно оценить"
    },
    {
      "title": "Диагноз",
      "content": "Основной диагноз по МКБ-10.\nСопутствующие."
    },
    {
      "title": "Назначения",
      "content": "- Медикаментозная терапия (с дозировками по весу/возрасту)\n- Режим, питание\n- Обследования\n- Дата повторного осмотра\n- Показания для экстренного обращения"
    }
  ],
  "summary": "Краткое резюме приёма в 1-2 предложения"
}

Правила:
- Педиатрическая терминология, учитывать возраст
- Заполняй ТОЛЬКО на основе данных
- Нет данных = "Данные не предоставлены"
- НЕ придумывай
- МКБ-10, русский язык""",
}


@app.get("/health")
async def health():
    return {"status": "ok", "version": "2.0.0"}


@app.post("/transcribe")
async def transcribe_audio(audio: UploadFile = File(...)):
    """Распознавание речи через Nexara API. Принимает аудиофайл любого формата."""
    if not NEXARA_API_KEY:
        raise HTTPException(status_code=500, detail="NEXARA_API_KEY не задан. Получите ключ на nexara.ru")

    # Проверяем формат файла
    filename = audio.filename or "audio.webm"
    ext = os.path.splitext(filename)[1].lower()
    if ext and ext not in ALLOWED_AUDIO:
        raise HTTPException(
            status_code=400,
            detail=f"Формат {ext} не поддерживается. Допустимые: MP3, WAV, M4A, OGG, FLAC, WebM",
        )

    # Сохраняем во временный файл
    suffix = ext or ".webm"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await audio.read()
        if len(content) > 100 * 1024 * 1024:  # 100 МБ лимит
            raise HTTPException(status_code=400, detail="Файл слишком большой (максимум 100 МБ)")
        tmp.write(content)
        tmp_path = tmp.name

    try:
        async with httpx.AsyncClient(timeout=600.0) as client:
            with open(tmp_path, "rb") as audio_file:
                response = await client.post(
                    NEXARA_API_URL,
                    headers={"Authorization": f"Bearer {NEXARA_API_KEY}"},
                    files={"file": (filename, audio_file)},
                    data={
                        "task": "transcribe",
                        "language": "ru",
                        "model": "whisper-1",
                        "response_format": "json",
                    },
                )

            if response.status_code != 200:
                err_text = response.text
                if "insufficient" in err_text.lower() or "quota" in err_text.lower():
                    raise HTTPException(status_code=429, detail="Недостаточно средств на аккаунте Nexara. Пополните баланс на nexara.ru")
                raise HTTPException(status_code=500, detail=f"Ошибка Nexara ({response.status_code}): {err_text[:200]}")

            data = response.json()
            text = data.get("text", "").strip()

            if not text:
                raise HTTPException(status_code=400, detail="Не удалось распознать речь в записи")

            return {"text": text, "filename": filename}

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Превышено время ожидания ответа от сервера распознавания. Попробуйте файл меньшего размера.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка распознавания речи: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/structure")
async def structure_text(
    text: str = Form(...),
    specialty: str = Form("psychiatrist"),
):
    """Структурирование текста через GigaChat API."""
    if not os.environ.get("GIGACHAT_AUTH_KEY"):
        raise HTTPException(status_code=500, detail="GIGACHAT_AUTH_KEY не задан")

    if specialty not in PROMPTS:
        raise HTTPException(status_code=400, detail=f"Неизвестная специальность: {specialty}")

    try:
        response_text = await gigachat_complete(
            messages=[
                {
                    "role": "user",
                    "content": f'{PROMPTS[specialty]}\n\nРасшифровка речи врача:\n"{text}"',
                }
            ],
            max_tokens=8192,
        )

        cleaned = response_text.strip()
        # Убираем markdown обёртки если есть
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[-1]
        if cleaned.endswith("```"):
            cleaned = cleaned.rsplit("```", 1)[0]
        cleaned = cleaned.strip()

        # Попытка 1: прямой парсинг
        try:
            result = json.loads(cleaned)
            return result
        except json.JSONDecodeError:
            pass

        # Попытка 2: убираем управляющие символы (кроме уже экранированных)
        def clean_json_string(s):
            # Заменяем literal newlines/tabs внутри строк на экранированные версии
            # Но не трогаем уже экранированные \n, \t
            result = []
            in_string = False
            escape = False
            for ch in s:
                if escape:
                    result.append(ch)
                    escape = False
                    continue
                if ch == '\\' and in_string:
                    result.append(ch)
                    escape = True
                    continue
                if ch == '"':
                    in_string = not in_string
                    result.append(ch)
                    continue
                if in_string:
                    if ch == '\n':
                        result.append('\\n')
                    elif ch == '\r':
                        result.append('\\r')
                    elif ch == '\t':
                        result.append('\\t')
                    elif ord(ch) < 0x20:
                        result.append(' ')
                    else:
                        result.append(ch)
                else:
                    result.append(ch)
            return ''.join(result)

        try:
            cleaned2 = clean_json_string(cleaned)
            result = json.loads(cleaned2)
            return result
        except json.JSONDecodeError:
            pass

        # Попытка 3: извлечь JSON из текста регулярным выражением
        json_match = re.search(r'\{[\s\S]*\}', cleaned)
        if json_match:
            try:
                cleaned3 = clean_json_string(json_match.group())
                result = json.loads(cleaned3)
                return result
            except json.JSONDecodeError:
                pass

        # Попытка 4: попросить Claude исправить
        raise json.JSONDecodeError("All parsing attempts failed", cleaned[:100], 0)

    except json.JSONDecodeError:
        try:
            fix_text = await gigachat_complete(
                messages=[
                    {"role": "user", "content": f"Этот JSON невалидный. Исправь его и верни ТОЛЬКО валидный JSON. Никаких пояснений, только JSON:\n\n{response_text[:12000]}"}
                ],
                max_tokens=8192,
            )
            fix_text = fix_text.strip()
            if fix_text.startswith("```"):
                fix_text = fix_text.split("\n", 1)[-1]
            if fix_text.endswith("```"):
                fix_text = fix_text.rsplit("```", 1)[0]
            fix_text = fix_text.strip()

            def clean_json_string2(s):
                result = []
                in_string = False
                escape = False
                for ch in s:
                    if escape:
                        result.append(ch)
                        escape = False
                        continue
                    if ch == '\\' and in_string:
                        result.append(ch)
                        escape = True
                        continue
                    if ch == '"':
                        in_string = not in_string
                        result.append(ch)
                        continue
                    if in_string and ch == '\n':
                        result.append('\\n')
                    elif in_string and ch == '\r':
                        result.append('\\r')
                    elif in_string and ch == '\t':
                        result.append('\\t')
                    elif in_string and ord(ch) < 0x20:
                        result.append(' ')
                    else:
                        result.append(ch)
                return ''.join(result)

            result = json.loads(clean_json_string2(fix_text))
            return result
        except Exception as fix_err:
            raise HTTPException(status_code=500, detail="Ошибка парсинга ответа. Попробуйте ещё раз.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка GigaChat: {str(e)}")


@app.post("/process")
async def process_audio(
    audio: UploadFile = File(...),
    specialty: str = Form("psychiatrist"),
):
    """Полный пайплайн: аудио → Whisper → Claude → документ."""
    transcription = await transcribe_audio(audio)
    text = transcription["text"]
    if not text.strip():
        raise HTTPException(status_code=400, detail="Не удалось распознать речь в записи")
    result = await structure_text(text=text, specialty=specialty)
    return {"transcript": text, "document": result}


# ─── Diagnostic Assistant ───

@app.post("/diagnose")
async def diagnose(
    sections: str = Form("[]"),
    patient_name: str = Form(""),
    transcript: str = Form(""),
):
    """Помощь с постановкой диагноза на основе данных приёма."""
    sections_data = json.loads(sections) if isinstance(sections, str) else sections
    sections_text = "\n".join([f"{s.get('title','')}: {s.get('content','')}" for s in sections_data if s.get('content') and s['content'] != 'Данные не предоставлены'])

    prompt = """Ты — опытный врач-психиатр, помогаешь коллеге сформулировать предварительный диагноз по МКБ-10.

ЗАДАЧА: на основе данных осмотра определи наиболее вероятный диагноз.

ТРЕБОВАНИЯ К ДИАГНОЗУ:
- Опирайся строго на критерии МКБ-10 и клинические рекомендации Минздрава РФ
- Выбирай наиболее точный код (с подрубриками: F20.0, F32.1, а не просто F20 или F32)
- При психотических симптомах (бред, галлюцинации, дезорганизация мышления) — рассматривай F20-F29 в первую очередь
- При аффективных нарушениях — F30-F39
- Диагноз формулируй: заболевание + синдром + степень тяжести + тип течения

КРИТИЧЕСКИ ВАЖНО — ФОРМАТ ОТВЕТА:
Все поля должны быть СТРОКАМИ (не объектами, не массивами, только текст).
Поле "treatment" — единый текст, НЕ JSON-объект.
Ответ — СТРОГО JSON без markdown и backticks:

{
  "diagnosis": "Полная формулировка диагноза строкой",
  "icd_code": "Код МКБ-10 с расшифровкой строкой, например: F20.0 Параноидная шизофрения",
  "justification": "Обоснование строкой: какие конкретные симптомы соответствуют каким критериям МКБ-10",
  "differential": "Дифференциальный диагноз строкой: 2-3 альтернативы с кодами и аргументами",
  "treatment": "Лечение строкой: 1) Фармакотерапия — препарат, доза начальная, доза целевая, длительность. 2) Психотерапия — методы. 3) Мониторинг.",
  "examinations": "Обследования строкой: что назначить и зачем"
}"""

    full_text = f"Данные пациента: {patient_name}\n\n{sections_text}"
    if transcript:
        full_text += f"\n\nИсходная расшифровка приёма:\n{transcript[:2000]}"

    try:
        message = await gigachat_complete(
            messages=[{"role": "user", "content": f"{prompt}\n\nДАННЫЕ ОСМОТРА:\n{full_text}"}],
            max_tokens=4096,
        )
        raw = message.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[-1]
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]
        raw = raw.strip()
        result = json.loads(raw)

        # Флатим вложенные объекты в строки (GigaChat иногда возвращает объекты вместо строк)
        def flatten_field(v) -> str:
            if v is None:
                return ""
            if isinstance(v, str):
                return v
            if isinstance(v, dict):
                parts = []
                for key, val in v.items():
                    parts.append(f"{key.capitalize()}: {flatten_field(val)}")
                return "\n\n".join(parts)
            if isinstance(v, list):
                return "\n".join(flatten_field(i) for i in v)
            return str(v)

        for field in ["diagnosis", "icd_code", "justification", "differential", "treatment", "examinations"]:
            if field in result:
                result[field] = flatten_field(result[field])

        return result
    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', raw)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                pass
        raise HTTPException(status_code=500, detail="Ошибка парсинга. Попробуйте ещё раз.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")


# ─── Structure by Template ───

@app.post("/structure-template")
async def structure_by_template(
    text: str = Form(...),
    template: UploadFile = File(...),
):
    """Структурирование по загруженному шаблону документа."""
    filename = template.filename or "template.txt"
    content_bytes = await template.read()

    ADMIN_KEYWORDS = [
        'зав. отделением', 'фамилия и.о', 'комиссия врачей',
        'консилиум', 'врачебная комиссия', 'лечащий врач'
    ]

    def parse_docx_runs(path: str) -> list:
        """
        Читает docx на уровне runs:
        - Bold run = начало новой секции (заголовок)
        - Non-bold run / параграф = подсказка к текущей секции
        Возвращает только медицинские секции (без подписей и оргчастей).
        """
        from docx import Document as D
        doc = D(path)

        sections = []
        current_title = None
        current_hints = []
        bold_buffer = []

        def flush_bold():
            nonlocal current_title, current_hints, bold_buffer
            if not bold_buffer:
                return
            raw = ' '.join(bold_buffer).strip().strip(':.')
            bold_buffer = []
            clean = re.sub(r'/строка[^/]*/', '', raw)
            clean = re.sub(r'\([^)]*\)', '', clean).strip().strip(':.')
            if not clean or len(clean) < 4:
                return
            if current_title:
                sections.append({
                    'title': current_title,
                    'hint': ' '.join(current_hints)
                })
            current_title = clean
            current_hints = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            for run in para.runs:
                t = run.text
                if not t.strip():
                    continue
                if run.bold:
                    bold_buffer.append(t.strip())
                else:
                    flush_bold()
                    ht = t.strip()
                    if ht and current_title:
                        current_hints.append(ht)

        flush_bold()
        if current_title:
            sections.append({'title': current_title, 'hint': ' '.join(current_hints)})

        # Убираем административные секции (подписи, комиссии)
        def is_admin(title):
            tl = title.lower()
            return any(k in tl for k in ADMIN_KEYWORDS)

        return [s for s in sections if not is_admin(s['title'])]

    def parse_text_fallback(content: str) -> list:
        """Fallback для .txt файлов — ищет по ключевым словам."""
        HEADER_KW = [
            "психический статус", "неврологическое", "соматическое",
            "назначения", "жалобы", "в дополнение", "по докладу",
            "протокол осмотра", "сон", "аппетит"
        ]
        sections = []
        current_title = None
        current_lines = []

        def flush():
            if current_title:
                sections.append({'title': current_title, 'hint': ' '.join(current_lines)})

        for line in content.replace("\r\n", "\n").split("\n"):
            s = line.strip()
            if not s:
                continue
            low = s.lower()
            if (len(s) < 120 and "нужное выбрать" not in low
                    and "/строка" not in low
                    and any(kw in low for kw in HEADER_KW)):
                flush()
                current_title = s.rstrip(":")
                current_lines = []
            else:
                current_lines.append(s)
        flush()
        return sections

    # ── Получаем секции ──
    parsed = []
    tmp_path = None
    try:
        if filename.endswith(".docx") or filename.endswith(".doc"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            parsed = parse_docx_runs(tmp_path)
        else:
            content_text = content_bytes.decode("utf-8", errors="ignore")
            parsed = parse_text_fallback(content_text)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    if not parsed:
        raise HTTPException(status_code=400,
                            detail="Не удалось извлечь структуру из шаблона. Проверьте файл.")

    # ── Скелет JSON ──
    skeleton = {"sections": [{"title": s["title"], "content": ""} for s in parsed]}
    skeleton_json = json.dumps(skeleton, ensure_ascii=False, indent=2)

    # ── Инструкции по каждой секции ──
    hints = []
    for s in parsed:
        hint = s["hint"]
        low = hint.lower()
        if "нужное выбрать" in low:
            opts = re.sub(r'/строка[^/]*/', '', hint)
            opts = re.sub(r'\([^)]*\)', lambda m: m.group() if 'нужное' in m.group() else '', opts)
            opts = re.sub(r'\(нужное выбрать[^)]*\)', '', opts).strip()[:400]
            hints.append(f'• {s["title"]}: ВЫБЕРИ подходящее из вариантов → {opts}')
        elif "/строка" in low or len(hint) < 30:
            hints.append(f'• {s["title"]}: заполни данными пациента, если нет — "не предъявляет"')
        else:
            hints.append(f'• {s["title"]}: заполни по смыслу из данных пациента')

    prompt = f"""Ты — врач-психиатр. Заполни разделы медицинского документа.

ДАННЫЕ ПАЦИЕНТА:
{text}

КАК ЗАПОЛНЯТЬ КАЖДЫЙ РАЗДЕЛ:
{chr(10).join(hints)}

СТРОГИЕ ПРАВИЛА:
1. Заполни ТОЛЬКО поле "content" в каждом разделе
2. НЕ добавляй новые разделы — их ровно {len(parsed)} штук
3. НЕ меняй поле "title"
4. Отвечай ТОЛЬКО JSON, без пояснений, без ```

{skeleton_json}"""

    try:
        message = await gigachat_complete(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=8192,
        )
        raw = message.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[-1]
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]
        raw = raw.strip()

        filled = json.loads(raw)
        sections_out = filled.get("sections", []) if isinstance(filled, dict) else []

        # Фильтр: только секции из скелета
        skeleton_titles = {s["title"] for s in parsed}
        sections_out = [s for s in sections_out if s.get("title") in skeleton_titles]

        # Добираем пропущенные
        filled_titles = {s["title"] for s in sections_out}
        for s in parsed:
            if s["title"] not in filled_titles:
                sections_out.append({"title": s["title"], "content": "не предъявляет"})

        # Восстанавливаем порядок
        order = {s["title"]: i for i, s in enumerate(parsed)}
        sections_out.sort(key=lambda s: order.get(s["title"], 999))

        return {
            "patient_name": filled.get("patient_name", "") if isinstance(filled, dict) else "",
            "diagnosis_code": filled.get("diagnosis_code", "") if isinstance(filled, dict) else "",
            "specialty": "psychiatrist",
            "sections": sections_out,
            "summary": filled.get("summary", "") if isinstance(filled, dict) else "",
        }

    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', raw)
        if match:
            try:
                filled = json.loads(match.group())
                return {
                    "patient_name": filled.get("patient_name", ""),
                    "diagnosis_code": filled.get("diagnosis_code", ""),
                    "specialty": "psychiatrist",
                    "sections": filled.get("sections", [{"title": s["title"], "content": "не предъявляет"} for s in parsed]),
                    "summary": filled.get("summary", ""),
                }
            except json.JSONDecodeError:
                pass
        raise HTTPException(status_code=500, detail="Ошибка парсинга ответа. Попробуйте ещё раз.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")


        for s in parsed:
            if s["title"] not in filled_titles:
                sections_out.append({"title": s["title"], "content": "не предъявляет"})

        # Восстанавливаем порядок как в шаблоне
        order = {s["title"]: i for i, s in enumerate(parsed)}
        sections_out.sort(key=lambda s: order.get(s["title"], 999))

        return {
            "patient_name": filled.get("patient_name", "") if isinstance(filled, dict) else "",
            "diagnosis_code": filled.get("diagnosis_code", "") if isinstance(filled, dict) else "",
            "specialty": "psychiatrist",
            "sections": sections_out,
            "summary": filled.get("summary", "") if isinstance(filled, dict) else "",
        }

    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', raw)
        if match:
            try:
                filled = json.loads(match.group())
                return {
                    "patient_name": filled.get("patient_name", ""),
                    "diagnosis_code": filled.get("diagnosis_code", ""),
                    "specialty": "psychiatrist",
                    "sections": filled.get("sections", [{"title": s["title"], "content": "не предъявляет"} for s in parsed]),
                    "summary": filled.get("summary", ""),
                }
            except json.JSONDecodeError:
                pass
        raise HTTPException(status_code=500, detail="Ошибка парсинга ответа. Попробуйте ещё раз.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")


@app.post("/records")
async def save_record(
    patient_name: str = Form(""),
    diagnosis_code: str = Form(""),
    specialty: str = Form(""),
    summary: str = Form(""),
    sections: str = Form("[]"),
    transcript: str = Form(""),
    authorization: str = Header(None),
):
    """Сохранить запись приёма в базу данных."""
    user = require_auth(authorization)
    record_id = str(uuid.uuid4())[:8]
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    conn = get_db()
    conn.execute(
        "INSERT INTO records (id, user_id, patient_name, diagnosis_code, specialty, summary, sections, transcript, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (record_id, user["id"], patient_name, diagnosis_code, specialty, summary, sections, transcript, now),
    )
    conn.commit()
    conn.close()
    return {"id": record_id, "created_at": now}


@app.get("/records")
async def list_records(authorization: str = Header(None)):
    """Список записей текущего врача."""
    user = require_auth(authorization)
    conn = get_db()
    rows = conn.execute(
        "SELECT id, patient_name, diagnosis_code, specialty, summary, created_at FROM records WHERE user_id = ? ORDER BY created_at DESC",
        (user["id"],),
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/records/{record_id}")
async def get_record(record_id: str, authorization: str = Header(None)):
    """Получить полную запись по ID."""
    user = require_auth(authorization)
    conn = get_db()
    row = conn.execute("SELECT * FROM records WHERE id = ? AND user_id = ?", (record_id, user["id"])).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Запись не найдена")
    record = dict(row)
    record["sections"] = json.loads(record["sections"])
    return record


@app.delete("/records/{record_id}")
async def delete_record(record_id: str, authorization: str = Header(None)):
    """Удалить запись."""
    user = require_auth(authorization)
    conn = get_db()
    conn.execute("DELETE FROM records WHERE id = ? AND user_id = ?", (record_id, user["id"]))
    conn.commit()
    conn.close()
    return {"deleted": record_id}


@app.patch("/records/{record_id}/diary")
async def append_diary_entry(
    record_id: str,
    sections: str = Form("[]"),
    transcript: str = Form(""),
    summary: str = Form(""),
    authorization: str = Header(None),
):
    """Добавить дневниковую запись к существующему пациенту."""
    user = require_auth(authorization)
    conn = get_db()
    row = conn.execute("SELECT * FROM records WHERE id = ? AND user_id = ?", (record_id, user["id"])).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Запись не найдена")

    existing = json.loads(row["sections"])
    new_entries = json.loads(sections)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    divider = {"title": f"── Дневник {now} ──", "content": "", "isDivider": True}
    combined = existing + [divider] + new_entries

    new_summary = row["summary"]
    if summary:
        new_summary = (row["summary"] or "") + f" | {now}: {summary}"

    conn.execute(
        "UPDATE records SET sections = ?, summary = ? WHERE id = ? AND user_id = ?",
        (json.dumps(combined, ensure_ascii=False), new_summary, record_id, user["id"]),
    )
    conn.commit()
    conn.close()
    return {"ok": True, "entries_added": len(new_entries)}


# ─── Word Export ───

from fastapi.responses import StreamingResponse
import io

@app.post("/export-word")
async def export_word(
    patient_name: str = Form(""),
    diagnosis_code: str = Form(""),
    specialty: str = Form(""),
    summary: str = Form(""),
    sections: str = Form("[]"),
):
    """Экспорт документа приёма в Word (.docx) по формату ПНД."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Первичный осмотр")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # Patient name
    if patient_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(patient_name)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'

    # Preamble
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_before = Pt(8)
    run = preamble.add_run(
        'Представился(лась) психиатром, разъяснены права согласно закону РФ '
        '"О психиатрической помощи и гарантиях прав граждан при ее оказании". '
        'На беседу согласен(а). Подтвердил(а) согласие на осмотр и/или лечение '
        'в письменной форме.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Sections
    sections_data = json.loads(sections) if isinstance(sections, str) else sections

    for sec in sections_data:
        title_text = sec.get("title", "")
        content_text = sec.get("content", "")

        if not content_text or content_text == "Данные не предоставлены":
            continue

        # Section title as bold inline prefix (like in the sample records)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run_title = p.add_run(f"{title_text}: ")
        run_title.bold = True
        run_title.font.name = 'Times New Roman'
        run_title.font.size = Pt(12)

        run_content = p.add_run(content_text)
        run_content.font.name = 'Times New Roman'
        run_content.font.size = Pt(12)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"PO_{patient_name.split()[0] if patient_name else 'patient'}.docx"

    from urllib.parse import quote
    encoded_filename = quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )


# ─── Serve React frontend in production ───

@app.on_event("startup")
async def mount_static():
    if STATIC_DIR.exists():
        app.mount("/static", StaticFiles(directory=str(STATIC_DIR / "static")), name="static-assets")


@app.get("/{full_path:path}")
async def serve_react(full_path: str):
    """Serve React app for all non-API routes."""
    if STATIC_DIR.exists():
        # Try to serve the exact file
        file_path = STATIC_DIR / full_path
        if file_path.is_file():
            return FileResponse(str(file_path))
        # Fallback to index.html for React routing
        index = STATIC_DIR / "index.html"
        if index.exists():
            return FileResponse(str(index))
    return {"detail": "Frontend not built. Run: cd frontend && npm run build"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
